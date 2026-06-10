from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as PdfImage,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "manual-assets"
OUT = ROOT / "manual-output"
DOCX_PATH = OUT / "Manual_Local_Growth_OS.docx"
PDF_PATH = OUT / "Manual_Local_Growth_OS.pdf"

ASSETS.mkdir(exist_ok=True)
OUT.mkdir(exist_ok=True)

INK = (23, 32, 28)
MOSS = (63, 111, 88)
CORAL = (217, 92, 74)
MAIZE = (232, 180, 77)
GLACIER = (74, 144, 164)
MIST = (238, 242, 239)
PAPER = (247, 245, 240)
STONE = (96, 90, 84)
WHITE = (255, 255, 255)


def font(size, bold=False):
    base = "arialbd.ttf" if bold else "arial.ttf"
    try:
        return ImageFont.truetype(base, size)
    except OSError:
        return ImageFont.load_default()


def rounded(draw, xy, fill, outline=None, width=1, radius=18):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def text(draw, xy, value, size=24, color=INK, bold=False, anchor=None):
    draw.text(xy, value, font=font(size, bold), fill=color, anchor=anchor)


def wrapped_text(draw, xy, value, max_chars=34, size=22, color=INK, bold=False, line_gap=6):
    words = value.split()
    lines = []
    current = ""
    for word in words:
        trial = f"{current} {word}".strip()
        if len(trial) > max_chars and current:
            lines.append(current)
            current = word
        else:
            current = trial
    if current:
        lines.append(current)
    x, y = xy
    for line in lines:
        text(draw, (x, y), line, size, color, bold)
        y += size + line_gap


def arrow(draw, start, end, color=MOSS, width=4):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    if x2 >= x1:
        pts = [(x2, y2), (x2 - 16, y2 - 8), (x2 - 16, y2 + 8)]
    else:
        pts = [(x2, y2), (x2 + 16, y2 - 8), (x2 + 16, y2 + 8)]
    draw.polygon(pts, fill=color)


def save_dashboard_map():
    img = Image.new("RGB", (1600, 900), PAPER)
    d = ImageDraw.Draw(img)
    text(d, (60, 48), "Dashboard: o cockpit comercial", 42, INK, True)
    text(d, (60, 100), "A tela responde: quem priorizar, o que enviar e se a prospecção está funcionando.", 24, STONE)

    labels = [
        ("Leads qualifiés", "24", "Base de salões mapeados", MOSS),
        ("Diagnostics prêts", "7", "Análises prontas para abordagem", MAIZE),
        ("Propositions", "4", "Ofertas em andamento", GLACIER),
        ("Taux réponse", "31%", "Sinal de tração da prospecção", CORAL),
    ]
    x = 60
    for title, value, desc, color in labels:
        rounded(d, (x, 170, x + 330, 330), WHITE, (220, 216, 210), radius=16)
        text(d, (x + 24, 198), title, 24, INK)
        text(d, (x + 24, 240), value, 48, (0, 0, 0), True)
        text(d, (x + 24, 295), desc, 20, STONE)
        d.rectangle((x, 170, x + 8, 330), fill=color)
        x += 370

    rounded(d, (60, 410, 930, 820), WHITE, (220, 216, 210), radius=18)
    text(d, (90, 445), "Leads prioritaires", 32, INK, True)
    text(d, (90, 490), "Mostra salões que merecem ação primeiro.", 22, STONE)
    lead_rows = [
        ("Atelier Coupe Lausanne", "Score 86/100", "Prioritaire", CORAL),
        ("Salon Belle Rive", "Score 74/100", "Diagnostic", MAIZE),
        ("Studio Hair & Care", "Score 68/100", "À contacter", GLACIER),
    ]
    y = 545
    for name, score, status, color in lead_rows:
        rounded(d, (90, y, 900, y + 75), (252, 252, 251), (232, 228, 222), radius=14)
        text(d, (115, y + 18), name, 23, INK, True)
        text(d, (115, y + 48), score, 18, STONE)
        rounded(d, (735, y + 18, 875, y + 50), (245, 239, 232), color, radius=16)
        text(d, (805, y + 34), status, 16, color, True, "mm")
        y += 92

    rounded(d, (990, 410, 1540, 670), WHITE, (220, 216, 210), radius=18)
    text(d, (1020, 445), "Prochaine proposition", 32, INK, True)
    text(d, (1020, 490), "Indica a proposta que precisa de atenção agora.", 22, STONE)
    rounded(d, (1020, 545, 1510, 625), MIST, None, radius=14)
    text(d, (1045, 565), "Pack Visibilité Locale", 24, INK, True)
    text(d, (1045, 600), "CHF 1'200 • Brouillon • À finaliser", 20, CORAL)

    arrow(d, (1260, 670), (1260, 780), CORAL)
    text(d, (1040, 805), "Ação esperada: finalizar a proposta.", 24, INK, True)
    img.save(ASSETS / "dashboard_explicado.png")


def save_workflow():
    img = Image.new("RGB", (1600, 760), WHITE)
    d = ImageDraw.Draw(img)
    text(d, (60, 50), "Fluxo de uso do MVP", 42, INK, True)
    steps = [
        ("1", "Mapear salão", "Coiffures em Lausanne"),
        ("2", "Qualificar lead", "Score, bairro, site, IG"),
        ("3", "Gerar diagnóstico", "Forças, riscos e ações"),
        ("4", "Enviar mensagem", "Contato manual"),
        ("5", "Criar proposta", "Oferta simples em CHF"),
        ("6", "Medir tração", "Resposta e rendez-vous"),
    ]
    x = 60
    y = 220
    for i, (num, title, desc) in enumerate(steps):
        rounded(d, (x, y, x + 220, y + 210), PAPER, (220, 216, 210), radius=22)
        d.ellipse((x + 20, y + 22, x + 66, y + 68), fill=MOSS)
        text(d, (x + 43, y + 45), num, 22, WHITE, True, "mm")
        text(d, (x + 20, y + 92), title, 24, INK, True)
        wrapped_text(d, (x + 20, y + 132), desc, max_chars=19, size=17, color=STONE)
        if i < len(steps) - 1:
            arrow(d, (x + 230, y + 105), (x + 285, y + 105), MOSS, 5)
        x += 255
    text(d, (60, 610), "Objetivo: vender serviços simples e aprender com conversas reais antes de construir uma plataforma maior.", 25, INK, True)
    img.save(ASSETS / "fluxo_mvp.png")


def save_lead_card():
    img = Image.new("RGB", (1600, 900), PAPER)
    d = ImageDraw.Draw(img)
    text(d, (60, 52), "Como ler um LeadCard", 42, INK, True)
    rounded(d, (150, 170, 1020, 725), WHITE, (218, 214, 207), radius=20)
    text(d, (200, 215), "Atelier Coupe Lausanne", 34, INK, True)
    rounded(d, (790, 205, 965, 250), (252, 231, 228), None, radius=22)
    text(d, (878, 228), "Prioritaire", 20, CORAL, True, "mm")
    text(d, (200, 275), "Flon", 24, STONE)
    rows = [("Score", "86/100"), ("Web", "atelier-coupe.ch"), ("IG", "@ateliercoupe_lsn")]
    y = 340
    for label, value in rows:
        text(d, (200, y), label, 24, INK, True)
        text(d, (340, y), value, 24, STONE)
        y += 58
    text(d, (200, 535), "Bon avis Google, faible cadence de posts,", 24, STONE)
    text(d, (200, 570), "réservation peu visible.", 24, STONE)
    rounded(d, (200, 630, 455, 680), MIST, None, radius=22)
    text(d, (328, 655), "Optimiser fiche Google", 20, INK, False, "mm")
    rounded(d, (480, 630, 735, 680), MIST, None, radius=22)
    text(d, (608, 655), "Assistant FAQ WhatsApp", 20, INK, False, "mm")

    notes = [
        ((1095, 215), (960, 225), "Status: etapa comercial atual."),
        ((1095, 345), (430, 355), "Score: prioridade relativa."),
        ((1095, 500), (630, 545), "Nota: leitura rápida para orientar a abordagem."),
        ((1095, 655), (610, 655), "Tags: ideias de serviços para vender."),
    ]
    for (tx, ty), end, label in notes:
        arrow(d, (tx - 30, ty + 10), end, GLACIER, 4)
        wrapped_text(d, (tx, ty), label, max_chars=32, size=23, color=INK, bold=True)
    img.save(ASSETS / "leadcard_explicado.png")


def save_diagnostic():
    img = Image.new("RGB", (1600, 900), WHITE)
    d = ImageDraw.Draw(img)
    text(d, (60, 52), "Nouveau diagnostic: venda consultiva", 42, INK, True)
    text(d, (60, 105), "A página transforma sinais simples do lead em uma conversa comercial em francês.", 24, STONE)

    sections = [
        ("Forces", "O que o salão já faz bem", ["Presença local identificável", "Boa base para conteúdo visual", "Sinais de reputação"]),
        ("Risques", "O que pode travar novos rendez-vous", ["Reserva pouco visível", "Provas sociais dispersas", "Pouco follow-up"]),
        ("Actions", "O que João pode propor", ["Optimiser fiche Google", "Scripts de relance", "Pack contenu + avis"]),
    ]
    x = 80
    for title, desc, bullets in sections:
        rounded(d, (x, 230, x + 460, 690), PAPER, (220, 216, 210), radius=20)
        text(d, (x + 30, 270), title, 34, MOSS, True)
        text(d, (x + 30, 320), desc, 22, STONE)
        y = 395
        for bullet in bullets:
            d.ellipse((x + 35, y + 8, x + 49, y + 22), fill=MAIZE)
            text(d, (x + 65, y), bullet, 24, INK)
            y += 72
        x += 510
    rounded(d, (210, 745, 1390, 825), MIST, None, radius=22)
    text(d, (800, 785), "Uso prático: enviar uma mensagem menos genérica e mais baseada em oportunidades reais.", 25, INK, True, "mm")
    img.save(ASSETS / "diagnostic_explicado.png")


def save_pages_map():
    img = Image.new("RGB", (1600, 900), PAPER)
    d = ImageDraw.Draw(img)
    text(d, (60, 50), "Mapa das páginas", 42, INK, True)
    pages = [
        ("Tableau", "Visão geral do pipeline e próximas ações.", MOSS),
        ("Leads", "Lista de salões mapeados e oportunidades.", GLACIER),
        ("Diagnostic", "Diagnóstico simulado em francês.", MAIZE),
        ("Propositions", "Ofertas comerciais simples.", CORAL),
        ("Messages", "Templates para contato manual.", MOSS),
        ("Rapports", "Métricas de validação comercial.", GLACIER),
    ]
    for i, (name, desc, color) in enumerate(pages):
        row = i // 2
        col = i % 2
        x = 80 + col * 760
        y = 155 + row * 230
        rounded(d, (x, y, x + 680, y + 160), WHITE, (220, 216, 210), radius=18)
        d.rectangle((x, y, x + 12, y + 160), fill=color)
        text(d, (x + 36, y + 38), name, 32, INK, True)
        text(d, (x + 36, y + 92), desc, 24, STONE)
    img.save(ASSETS / "mapa_paginas.png")


def save_metrics_ladder():
    img = Image.new("RGB", (1600, 760), WHITE)
    d = ImageDraw.Draw(img)
    text(d, (60, 50), "Como interpretar as métricas", 42, INK, True)
    items = [
        ("Leads", "Tenho salões suficientes no radar?"),
        ("Diagnostics", "Tenho material concreto para abrir conversa?"),
        ("Messages", "Estou abordando de forma consistente?"),
        ("Réponses", "O mercado está reagindo?"),
        ("Propositions", "As conversas viram oportunidade comercial?"),
        ("Rendez-vous", "A venda está saindo do papel?"),
    ]
    x = 120
    y = 210
    for idx, (label, question) in enumerate(items):
        rounded(d, (x, y, x + 1320, y + 70), MIST if idx % 2 == 0 else PAPER, None, radius=16)
        text(d, (x + 30, y + 20), label, 25, INK, True)
        text(d, (x + 300, y + 20), question, 25, STONE)
        y += 85
    img.save(ASSETS / "metricas_explicadas.png")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="DADCE0"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:color"), color)
        borders.append(tag)


def add_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.75)
    table.columns[1].width = Inches(4.55)
    hdr = table.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Sentido"
    for cell in hdr:
        set_cell_shading(cell, "E8EEF5")
        set_cell_border(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
    for label, meaning in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = meaning
        for cell in cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Manual do Local Growth OS")
    run.font.name = "Calibri"
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(23, 32, 28)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Guia visual para entender o MVP de vendas IA + marketing local")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(85, 85, 85)

    doc.add_paragraph()
    doc.add_picture(str(ASSETS / "dashboard_explicado.png"), width=Inches(6.5))
    cap = doc.add_paragraph("Figura 1 - Leitura geral do Dashboard.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(85, 85, 85)
    doc.add_page_break()


def add_image(doc, filename, caption):
    doc.add_picture(str(ASSETS / filename), width=Inches(6.5))
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(85, 85, 85)


def build_doc():
    doc = Document()
    style_doc(doc)
    add_title(doc)

    doc.add_heading("1. O que é o MVP", level=1)
    doc.add_paragraph(
        "O Local Growth OS é um cockpit comercial para João Pedro validar a venda de serviços de IA + marketing "
        "para pequenos negócios da Suisse romande. O primeiro nicho é salões de coiffure em Lausanne."
    )
    doc.add_paragraph(
        "Ele ainda não é uma plataforma completa. A função é organizar prospecção, diagnósticos, propostas e métricas "
        "mínimas antes de investir em integrações, banco de dados ou automação."
    )
    add_table(
        doc,
        [
            ("Não é CRM completo", "É um painel de validação comercial com dados mockados."),
            ("Não envia mensagens", "Os templates servem para copiar, adaptar e enviar manualmente."),
            ("Não usa APIs externas", "A fase atual evita complexidade até provar demanda."),
            ("Não confirma receita", "Pipeline e propostas representam oportunidade, não dinheiro recebido."),
        ],
    )

    doc.add_heading("2. Fluxo de trabalho", level=1)
    add_image(doc, "fluxo_mvp.png", "Figura 2 - Fluxo comercial recomendado para usar o MVP.")
    doc.add_paragraph(
        "A lógica é simples: mapear salões, qualificar os melhores, preparar um diagnóstico em francês, abordar "
        "manualmente, criar uma proposta simples e medir se o mercado responde."
    )

    doc.add_heading("3. Como entender o Dashboard", level=1)
    doc.add_paragraph(
        "O Dashboard é a tela do seu print. Ele existe para responder rapidamente onde está a energia comercial: "
        "quantos leads existem, quantos diagnósticos podem ser usados, quantas propostas estão abertas e se a "
        "prospecção está gerando resposta."
    )
    add_table(
        doc,
        [
            ("Leads qualifiés", "Quantidade de salões minimamente interessantes no radar."),
            ("Diagnostics prêts", "Diagnósticos preparados para abrir conversa comercial."),
            ("Propositions", "Propostas em andamento e valor potencial do pipeline."),
            ("Taux réponse", "Percentual de respostas obtidas em mensagens manuais."),
        ],
    )
    doc.add_paragraph(
        "Exemplo: 24 leads qualificados não significa 24 clientes. Significa 24 salões que parecem ter potencial. "
        "31% de resposta significa que a abordagem manual está recebendo algum retorno."
    )

    doc.add_heading("4. Como ler um LeadCard", level=1)
    add_image(doc, "leadcard_explicado.png", "Figura 3 - Anatomia de um card de lead.")
    doc.add_paragraph(
        "O card do lead resume a oportunidade comercial. O score ajuda a priorizar, o status mostra a etapa, "
        "a nota explica o problema visível e as tags sugerem serviços que João pode vender."
    )

    doc.add_heading("5. Página Nouveau diagnostic", level=1)
    add_image(doc, "diagnostic_explicado.png", "Figura 4 - Estrutura do diagnóstico simulado.")
    doc.add_paragraph(
        "O diagnóstico transforma sinais simples em uma abordagem consultiva. Em vez de dizer apenas que João faz "
        "marketing e IA, a conversa começa com oportunidades concretas: reserva pouco visível, Google Business, "
        "follow-up e mensagens melhores."
    )

    doc.add_heading("6. Mapa das páginas", level=1)
    add_image(doc, "mapa_paginas.png", "Figura 5 - Papel de cada área do app.")
    doc.add_paragraph(
        "Cada página tem uma função específica. Tableau mostra prioridades, Leads mostra a base, Diagnostic prepara "
        "a conversa, Propositions organiza ofertas, Messages apoia a prospecção manual e Rapports mede validação."
    )

    doc.add_heading("7. Interpretação das métricas", level=1)
    add_image(doc, "metricas_explicadas.png", "Figura 6 - Métricas como funil de validação.")
    doc.add_paragraph(
        "As métricas devem ser lidas como perguntas de validação. Se os leads crescem mas as respostas não aparecem, "
        "a mensagem talvez esteja fraca. Se respostas aparecem mas propostas não avançam, a oferta precisa ficar mais clara."
    )

    doc.add_heading("8. Rotina prática para João", level=1)
    doc.add_paragraph("Rotina diária sugerida:")
    for item in [
        "Abrir o Dashboard e escolher 3 leads prioritários.",
        "Ler o LeadCard e identificar a oportunidade principal.",
        "Abrir Diagnostic e preparar a abordagem em francês.",
        "Usar um template em Messages e enviar manualmente.",
        "Registrar respostas e mover mentalmente o lead para proposta ou follow-up.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph("Rotina semanal sugerida:")
    for item in [
        "Contar leads qualificados novos.",
        "Medir diagnósticos preparados e enviados.",
        "Calcular taxa real de resposta.",
        "Revisar propostas abertas e relances pendentes.",
        "Ajustar oferta, preço ou mensagem com base nas conversas.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("9. Limites intencionais", level=1)
    doc.add_paragraph(
        "O MVP não tem login, banco de dados, cadastro real pela interface, geração real via IA, WhatsApp, Google Maps, "
        "pagamento ou envio automático. Esses limites são intencionais: primeiro validar venda, depois automatizar."
    )

    doc.add_heading("10. Próximos passos", level=1)
    for item in [
        "Testar a abordagem com 10 a 20 salões reais em Lausanne.",
        "Substituir mocks por planilha ou banco simples.",
        "Adicionar edição de leads e status.",
        "Gerar diagnóstico real com IA quando a proposta estiver validada.",
        "Exportar proposta em PDF.",
        "Integrar mensagens somente depois de provar tração comercial.",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_section(WD_SECTION.CONTINUOUS)
    doc.save(DOCX_PATH)


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="ManualTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=28,
            leading=34,
            textColor=colors.HexColor("#17201C"),
            alignment=TA_CENTER,
            spaceAfter=12,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ManualSubtitle",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=13,
            leading=18,
            textColor=colors.HexColor("#555555"),
            alignment=TA_CENTER,
            spaceAfter=24,
        )
    )
    styles.add(
        ParagraphStyle(
            name="H1Custom",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=17,
            leading=22,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=18,
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BodyCustom",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=15,
            textColor=colors.HexColor("#17201C"),
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Caption",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=8.5,
            leading=11,
            textColor=colors.HexColor("#555555"),
            alignment=TA_CENTER,
            spaceBefore=4,
            spaceAfter=12,
        )
    )
    return styles


def pdf_img(filename, caption):
    path = ASSETS / filename
    with Image.open(path) as img:
        width_px, height_px = img.size
    max_w = 6.25 * inch
    max_h = 4.15 * inch
    ratio = min(max_w / width_px, max_h / height_px)
    draw_w = width_px * ratio
    draw_h = height_px * ratio
    return KeepTogether(
        [
            PdfImage(str(path), width=draw_w, height=draw_h),
            Paragraph(caption, pdf_styles()["Caption"]),
        ]
    )


def pdf_table(rows):
    data = [["Item", "Sentido"], *rows]
    table = Table(data, colWidths=[1.7 * inch, 4.6 * inch])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#17201C")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (0, -1), "Helvetica-Bold"),
                ("FONTNAME", (1, 1), (1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9.5),
                ("LEADING", (0, 0), (-1, -1), 12),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#DADCE0")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    return table


def bullet_list(items, styles):
    return ListFlowable(
        [ListItem(Paragraph(item, styles["BodyCustom"])) for item in items],
        bulletType="bullet",
        start="circle",
        leftIndent=18,
    )


def numbered_list(items, styles):
    return ListFlowable(
        [ListItem(Paragraph(item, styles["BodyCustom"])) for item in items],
        bulletType="1",
        leftIndent=18,
    )


def build_pdf():
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
        title="Manual do Local Growth OS",
        author="joaopedro.chat Growth Lab",
    )

    story = [
        Paragraph("Manual do Local Growth OS", styles["ManualTitle"]),
        Paragraph("Guia visual para entender o MVP de vendas IA + marketing local", styles["ManualSubtitle"]),
        pdf_img("dashboard_explicado.png", "Figura 1 - Leitura geral do Dashboard."),
        PageBreak(),
        Paragraph("1. O que é o MVP", styles["H1Custom"]),
        Paragraph(
            "O Local Growth OS é um cockpit comercial para João Pedro validar a venda de serviços de IA + marketing "
            "para pequenos negócios da Suisse romande. O primeiro nicho é salões de coiffure em Lausanne.",
            styles["BodyCustom"],
        ),
        Paragraph(
            "Ele ainda não é uma plataforma completa. A função é organizar prospecção, diagnósticos, propostas e métricas "
            "mínimas antes de investir em integrações, banco de dados ou automação.",
            styles["BodyCustom"],
        ),
        pdf_table(
            [
                ["Não é CRM completo", "É um painel de validação comercial com dados mockados."],
                ["Não envia mensagens", "Os templates servem para copiar, adaptar e enviar manualmente."],
                ["Não usa APIs externas", "A fase atual evita complexidade até provar demanda."],
                ["Não confirma receita", "Pipeline e propostas representam oportunidade, não dinheiro recebido."],
            ]
        ),
        Paragraph("2. Fluxo de trabalho", styles["H1Custom"]),
        pdf_img("fluxo_mvp.png", "Figura 2 - Fluxo comercial recomendado para usar o MVP."),
        Paragraph(
            "A lógica é simples: mapear salões, qualificar os melhores, preparar um diagnóstico em francês, abordar "
            "manualmente, criar uma proposta simples e medir se o mercado responde.",
            styles["BodyCustom"],
        ),
        PageBreak(),
        Paragraph("3. Como entender o Dashboard", styles["H1Custom"]),
        Paragraph(
            "O Dashboard é a tela do seu print. Ele existe para responder rapidamente onde está a energia comercial: "
            "quantos leads existem, quantos diagnósticos podem ser usados, quantas propostas estão abertas e se a "
            "prospecção está gerando resposta.",
            styles["BodyCustom"],
        ),
        pdf_table(
            [
                ["Leads qualifiés", "Quantidade de salões minimamente interessantes no radar."],
                ["Diagnostics prêts", "Diagnósticos preparados para abrir conversa comercial."],
                ["Propositions", "Propostas em andamento e valor potencial do pipeline."],
                ["Taux réponse", "Percentual de respostas obtidas em mensagens manuais."],
            ]
        ),
        Paragraph(
            "Exemplo: 24 leads qualificados não significa 24 clientes. Significa 24 salões que parecem ter potencial. "
            "31% de resposta significa que a abordagem manual está recebendo algum retorno.",
            styles["BodyCustom"],
        ),
        Paragraph("4. Como ler um LeadCard", styles["H1Custom"]),
        pdf_img("leadcard_explicado.png", "Figura 3 - Anatomia de um card de lead."),
        Paragraph(
            "O card do lead resume a oportunidade comercial. O score ajuda a priorizar, o status mostra a etapa, "
            "a nota explica o problema visível e as tags sugerem serviços que João pode vender.",
            styles["BodyCustom"],
        ),
        PageBreak(),
        Paragraph("5. Página Nouveau diagnostic", styles["H1Custom"]),
        pdf_img("diagnostic_explicado.png", "Figura 4 - Estrutura do diagnóstico simulado."),
        Paragraph(
            "O diagnóstico transforma sinais simples em uma abordagem consultiva. Em vez de dizer apenas que João faz "
            "marketing e IA, a conversa começa com oportunidades concretas: reserva pouco visível, Google Business, "
            "follow-up e mensagens melhores.",
            styles["BodyCustom"],
        ),
        Paragraph("6. Mapa das páginas", styles["H1Custom"]),
        pdf_img("mapa_paginas.png", "Figura 5 - Papel de cada área do app."),
        Paragraph(
            "Cada página tem uma função específica. Tableau mostra prioridades, Leads mostra a base, Diagnostic prepara "
            "a conversa, Propositions organiza ofertas, Messages apoia a prospecção manual e Rapports mede validação.",
            styles["BodyCustom"],
        ),
        PageBreak(),
        Paragraph("7. Interpretação das métricas", styles["H1Custom"]),
        pdf_img("metricas_explicadas.png", "Figura 6 - Métricas como funil de validação."),
        Paragraph(
            "As métricas devem ser lidas como perguntas de validação. Se os leads crescem mas as respostas não aparecem, "
            "a mensagem talvez esteja fraca. Se respostas aparecem mas propostas não avançam, a oferta precisa ficar mais clara.",
            styles["BodyCustom"],
        ),
        Paragraph("8. Rotina prática para João", styles["H1Custom"]),
        Paragraph("Rotina diária sugerida:", styles["BodyCustom"]),
        bullet_list(
            [
                "Abrir o Dashboard e escolher 3 leads prioritários.",
                "Ler o LeadCard e identificar a oportunidade principal.",
                "Abrir Diagnostic e preparar a abordagem em francês.",
                "Usar um template em Messages e enviar manualmente.",
                "Registrar respostas e mover mentalmente o lead para proposta ou follow-up.",
            ],
            styles,
        ),
        Paragraph("Rotina semanal sugerida:", styles["BodyCustom"]),
        bullet_list(
            [
                "Contar leads qualificados novos.",
                "Medir diagnósticos preparados e enviados.",
                "Calcular taxa real de resposta.",
                "Revisar propostas abertas e relances pendentes.",
                "Ajustar oferta, preço ou mensagem com base nas conversas.",
            ],
            styles,
        ),
        Paragraph("9. Limites intencionais", styles["H1Custom"]),
        Paragraph(
            "O MVP não tem login, banco de dados, cadastro real pela interface, geração real via IA, WhatsApp, Google Maps, "
            "pagamento ou envio automático. Esses limites são intencionais: primeiro validar venda, depois automatizar.",
            styles["BodyCustom"],
        ),
        Paragraph("10. Próximos passos", styles["H1Custom"]),
        numbered_list(
            [
                "Testar a abordagem com 10 a 20 salões reais em Lausanne.",
                "Substituir mocks por planilha ou banco simples.",
                "Adicionar edição de leads e status.",
                "Gerar diagnóstico real com IA quando a proposta estiver validada.",
                "Exportar proposta em PDF.",
                "Integrar mensagens somente depois de provar tração comercial.",
            ],
            styles,
        ),
    ]
    doc.build(story)


if __name__ == "__main__":
    save_dashboard_map()
    save_workflow()
    save_lead_card()
    save_diagnostic()
    save_pages_map()
    save_metrics_ladder()
    build_doc()
    build_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)
